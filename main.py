import streamlit as st
import re
import fitz  # PyMuPDF
import pandas as pd
from io import BytesIO
from docx import Document

# Título da aplicação
st.title("Analisador de Palavras-Chave em Relatórios PDF")
st.write("Faça upload de um arquivo PDF e procure por palavras-chave específicas.")

# Entrada de palavras-chave pelo usuário
user_input = st.text_input("Digite até 3 palavras separadas por vírgulas:")
palavras = [palavra.strip() for palavra in user_input.split(",")][:3] if user_input else []

# Upload do arquivo PDF
uploaded_file = st.file_uploader("Carregue um arquivo PDF para análise", type="pdf")

if st.button("Iniciar Busca") and uploaded_file and palavras:
    try:
        # Carregar o PDF usando PyMuPDF
        pdf_doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        results_text = []
        word_frequencies = {}

        # Processar cada página do PDF
        for page_num, page in enumerate(pdf_doc, start=1):
            text = page.get_text("text")

            for palavra in palavras:
                palavra_regex = re.compile(r'\b' + re.escape(palavra) + r'\b', re.IGNORECASE)
                if palavra not in word_frequencies:
                    word_frequencies[palavra] = 0

                # Encontrar todas as ocorrências da palavra na página
                for match in palavra_regex.finditer(text):
                    start = max(match.start() - 500, 0)
                    end = min(match.end() + 500, len(text))
                    context = text[start:end]

                    results_text.append({
                        'Arquivo': uploaded_file.name,
                        'Palavra': palavra,
                        'Página': page_num,
                        'Ocorrência': context
                    })

                    word_frequencies[palavra] += 1

        # Exibir os resultados
        if results_text:
            st.success("Busca concluída! Abaixo estão os resultados encontrados.")
            for result in results_text:
                st.write(f"**Arquivo**: {result['Arquivo']}")
                st.write(f"**Palavra**: {result['Palavra']}")
                st.write(f"**Página**: {result['Página']}")
                st.write(f"**Ocorrência**: {result['Ocorrência']}")
                st.write("---")

            # Salvar resultados em arquivos
            # Frequências em Excel
            df_freq = pd.DataFrame(list(word_frequencies.items()), columns=['Palavra', 'Frequência'])
            excel_buffer = BytesIO()
            df_freq.to_excel(excel_buffer, index=False)
            excel_buffer.seek(0)

            # Ocorrências em Word
            doc = Document()
            doc.add_heading('Ocorrências de Palavras-Chave', 0)
            for result in results_text:
                doc.add_paragraph(f"Arquivo: {result['Arquivo']}")
                doc.add_paragraph(f"Palavra: {result['Palavra']}")
                doc.add_paragraph(f"Página: {result['Página']}")
                doc.add_paragraph(f"Ocorrência: {result['Ocorrência']}")
                doc.add_paragraph("\n")

            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            # Botões para download
            st.download_button("Baixar Frequências (Excel)", excel_buffer, file_name="frequencias_palavras.xlsx")
            st.download_button("Baixar Ocorrências (Word)", word_buffer, file_name="ocorrencias_palavras.docx")
        else:
            st.warning("Nenhuma ocorrência encontrada.")
    except Exception as e:
        st.error(f"Ocorreu um erro durante a análise: {e}")
else:
    if not uploaded_file:
        st.warning("Por favor, carregue um arquivo PDF.")
    if not palavras:
        st.warning("Por favor, insira até 3 palavras para busca.")
